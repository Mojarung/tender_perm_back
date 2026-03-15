"""Document generation service — creates .docx NMCK justification (Приказ №567)."""

import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src.services.num_to_words_ru import number_to_words_ru

logger = logging.getLogger(__name__)

_FONT_NAME = "Times New Roman"
_FONT_SIZE = Pt(12)


# ── Helpers ──────────────────────────────────────────────────────────────────


def _format_price(value: float) -> str:
    """Format price Russian-style: '1 234,56'."""
    integer_part = int(value)
    decimal_part = round((value - integer_part) * 100)
    int_str = f"{integer_part:,}".replace(",", " ")
    return f"{int_str},{decimal_part:02d}"


def _format_date(val: Any) -> str:
    """Convert datetime / string / None to 'DD.MM.YYYY'."""
    if val is None:
        return "—"
    if isinstance(val, datetime):
        return val.strftime("%d.%m.%Y")
    s = str(val).strip()
    if not s:
        return "—"
    # Already formatted
    if len(s) == 10 and s[2] == "." and s[5] == ".":
        return s
    # ISO format
    for fmt in ("%Y-%m-%d", "%Y-%m-%dT%H:%M:%S", "%d.%m.%Y"):
        try:
            return datetime.strptime(s[:10], fmt).strftime("%d.%m.%Y")
        except ValueError:
            continue
    return s


def _set_cell_shading(cell, color_hex: str) -> None:
    """Set cell background colour via XML."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def _style_paragraph(paragraph, bold: bool = False, size: Pt = _FONT_SIZE,
                     alignment=None, space_after: Pt | None = None) -> None:
    """Apply common paragraph styling."""
    if alignment is not None:
        paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    if space_after is not None:
        fmt.space_after = space_after
    for run in paragraph.runs:
        run.font.name = _FONT_NAME
        run.font.size = size
        run.bold = bold


def _add_styled_paragraph(doc: Document, text: str, bold: bool = False,
                          size: Pt = _FONT_SIZE,
                          alignment=None,
                          space_after: Pt | None = None):
    """Add a paragraph with consistent styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = _FONT_NAME
    run.font.size = size
    run.bold = bold
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    return p


def _style_table(table, header_row: bool = True) -> None:
    """Apply Table Grid + header shading."""
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if header_row and table.rows:
        for cell in table.rows[0].cells:
            _set_cell_shading(cell, "D9D9D9")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = _FONT_NAME
                    run.font.size = _FONT_SIZE
                    run.bold = True


def _set_cell_text(cell, text: str, bold: bool = False,
                   alignment=None) -> None:
    """Set cell text with styling."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = _FONT_NAME
    run.font.size = _FONT_SIZE
    run.bold = bold
    if alignment is not None:
        p.alignment = alignment


# ── Main Generator ───────────────────────────────────────────────────────────


def generate_nmck_document(
    output_dir: Path,
    session_id: str,
    target_name: str,
    analogs: list[dict[str, Any]],
    prices: list[dict[str, Any]],
    outliers: list[dict[str, Any]],
    calculation: dict[str, Any],
    justification: list[dict[str, Any]],
    quantity: float = 1.0,
    region: str | None = None,
    unit: str | None = None,
) -> str:
    """Generate .docx NMCK justification per Приказ №567 (Приложение 1)."""
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"nmck_justification_{session_id}.docx"

    doc = Document()

    # ── Page margins: left 3cm, right 1.5cm, top 2cm, bottom 2cm ──
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # ── Default font ──
    style = doc.styles["Normal"]
    style.font.name = _FONT_NAME
    style.font.size = _FONT_SIZE

    today = datetime.now().strftime("%d.%m.%Y")

    # ── 1. Title ─────────────────────────────────────────────────────────
    _add_styled_paragraph(
        doc, "ОБОСНОВАНИЕ\nначальной (максимальной) цены контракта",
        bold=True, size=Pt(14),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(12),
    )
    _add_styled_paragraph(
        doc,
        f"Предмет контракта: {target_name}",
        space_after=Pt(12),
    )

    # ── 2. Characteristics ───────────────────────────────────────────────
    _add_styled_paragraph(doc, "1. Характеристики объекта закупки", bold=True,
                          space_after=Pt(6))
    first_analog = analogs[0] if analogs else {}
    char_name = first_analog.get("name", target_name)
    _add_styled_paragraph(doc, f"Наименование: {char_name}")
    attrs = first_analog.get("attributes", {})
    if attrs:
        for key, val in attrs.items():
            _add_styled_paragraph(doc, f"  {key}: {val}")
    if region:
        _add_styled_paragraph(doc, f"Регион поставки: {region}")
    if unit:
        _add_styled_paragraph(doc, f"Единица измерения: {unit}")
    _add_styled_paragraph(
        doc, f"Количество: {_format_price(quantity) if quantity != int(quantity) else int(quantity)}",
        space_after=Pt(12),
    )

    # ── 3. Method ────────────────────────────────────────────────────────
    _add_styled_paragraph(doc, "2. Метод определения НМЦК", bold=True,
                          space_after=Pt(6))
    _add_styled_paragraph(
        doc,
        "Начальная (максимальная) цена контракта определена "
        "методом сопоставимых рыночных цен (анализа рынка) "
        "в соответствии с ч. 6 ст. 22 Федерального закона от 05.04.2013 "
        "№ 44-ФЗ «О контрактной системе в сфере закупок товаров, работ, "
        "услуг для обеспечения государственных и муниципальных нужд» "
        "и Приказа Минэкономразвития России от 02.10.2013 № 567 "
        "«Об утверждении Методических рекомендаций по применению методов "
        "определения начальной (максимальной) цены контракта, цены "
        "контракта, заключаемого с единственным поставщиком "
        "(подрядчиком, исполнителем)».",
        space_after=Pt(12),
    )

    # ── 4. Sources ───────────────────────────────────────────────────────
    _add_styled_paragraph(doc, "3. Источники ценовой информации", bold=True,
                          space_after=Pt(6))
    for i, p in enumerate(prices, 1):
        contract_id = p.get("Идентификатор контракта", p.get("contract_id", ""))
        date = _format_date(p.get("Дата заключения контракта", p.get("contract_date")))
        rgn = p.get("Регион заказчика", p.get("region", ""))
        source = p.get("_source")
        if source == "manual":
            _add_styled_paragraph(doc, f"{i}. Ручной ввод цены — {rgn}")
        else:
            _add_styled_paragraph(
                doc, f"{i}. Контракт № {contract_id} от {date}, {rgn}")
    _add_styled_paragraph(doc, "", space_after=Pt(6))

    # ── 5. Formula ───────────────────────────────────────────────────────
    _add_styled_paragraph(doc, "4. Формула расчёта НМЦК", bold=True,
                          space_after=Pt(6))
    _add_styled_paragraph(
        doc,
        "НМЦК = v × (Σ цi) / n",
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(6),
    )
    _add_styled_paragraph(doc, "где:")
    _add_styled_paragraph(doc, "  v — количество товара (работ, услуг);")
    _add_styled_paragraph(doc, "  n — количество значений, используемых в расчёте;")
    _add_styled_paragraph(
        doc,
        "  цi — цена единицы товара (работы, услуги), приведённая "
        "с учётом коэффициента пересчёта.",
        space_after=Pt(12),
    )

    # ── 6. Calculation table ─────────────────────────────────────────────
    _add_styled_paragraph(doc, "5. Расчёт НМЦК", bold=True, space_after=Pt(6))

    n_prices = len(prices)
    cols = 5
    # header + price rows + 6 summary rows
    table = doc.add_table(rows=1 + n_prices + 6, cols=cols)
    _style_table(table)

    # Header
    headers = ["№", "Источник", "Цена за ед., руб.", "Врем. коэфф.", "Привед. цена, руб."]
    for j, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[j], h, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Price rows
    total_adjusted = 0.0
    for i, p in enumerate(prices):
        row = table.rows[1 + i]
        raw_price = float(p.get("Цена за единицу", p.get("price", 0)))
        tw = float(p.get("time_weight", 1.0))
        adjusted = raw_price * tw
        total_adjusted += adjusted

        source = p.get("_source")
        if source == "manual":
            src_text = "Ручной ввод"
        else:
            cid = p.get("Идентификатор контракта", p.get("contract_id", ""))
            cdate = _format_date(p.get("Дата заключения контракта", p.get("contract_date")))
            src_text = f"Контракт № {cid} от {cdate}"

        _set_cell_text(row.cells[0], str(i + 1), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row.cells[1], src_text)
        _set_cell_text(row.cells[2], _format_price(raw_price),
                       alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(row.cells[3], f"{tw:.4f}",
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row.cells[4], _format_price(adjusted),
                       alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # Summary rows
    avg_price = calculation.get("weighted_average_price", 0)
    cv = calculation.get("coefficient_of_variation", 0)
    nmck_unit = calculation.get("nmck_per_unit", 0)
    total_nmck = calculation.get("total_nmck", 0)

    summary_data = [
        ("Количество ценовых данных (n)", str(n_prices)),
        ("Средняя цена за единицу, руб.", _format_price(avg_price)),
        ("Коэффициент вариации (КВ), %", f"{cv:.1f}"),
        ("Количество (v)", str(int(quantity)) if quantity == int(quantity) else _format_price(quantity)),
        ("НМЦК за единицу, руб.", _format_price(nmck_unit)),
        ("ИТОГО НМЦК, руб.", _format_price(total_nmck)),
    ]
    for idx, (label, value) in enumerate(summary_data):
        r = table.rows[1 + n_prices + idx]
        # Merge first 4 cells for the label
        r.cells[0].merge(r.cells[3])
        _set_cell_text(r.cells[0], label, bold=True)
        _set_cell_text(r.cells[4], value, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    _add_styled_paragraph(doc, "", space_after=Pt(6))

    # ── 7. Homogeneity ──────────────────────────────────────────────────
    _add_styled_paragraph(doc, "6. Однородность ценовой информации", bold=True,
                          space_after=Pt(6))
    is_homo = calculation.get("is_homogeneous", True)
    verdict = "однородна" if is_homo else "неоднородна"
    _add_styled_paragraph(
        doc,
        f"Коэффициент вариации составляет {cv:.1f}%. "
        f"Совокупность ценовой информации {verdict}.",
    )
    if not is_homo:
        _add_styled_paragraph(
            doc,
            "Примечание: при неоднородности выборки (КВ > 33%) "
            "рекомендуется провести дополнительный анализ рынка "
            "или обосновать использование данной ценовой информации.",
        )
    _add_styled_paragraph(doc, "", space_after=Pt(6))

    # ── 8. Outliers ──────────────────────────────────────────────────────
    if outliers:
        _add_styled_paragraph(doc, "7. Исключённые ценовые данные", bold=True,
                              space_after=Pt(6))
        _add_styled_paragraph(
            doc,
            f"Из расчёта исключено {len(outliers)} позиций "
            "как статистически выбивающихся из выборки:",
        )
        for o in outliers:
            price_val = float(o.get("Цена за единицу", o.get("price", 0)))
            name_val = o.get("Наименование позиции СТЕ", o.get("cte_name", ""))
            _add_styled_paragraph(
                doc, f"  — {name_val}: {_format_price(price_val)} руб.")
        _add_styled_paragraph(doc, "", space_after=Pt(6))

    # ── 9. Total ─────────────────────────────────────────────────────────
    section_num = 8 if outliers else 7
    _add_styled_paragraph(
        doc, f"{section_num}. Итого", bold=True, space_after=Pt(6))
    words = number_to_words_ru(total_nmck)
    _add_styled_paragraph(
        doc,
        f"НМЦК составляет: {_format_price(total_nmck)} руб. ({words})",
        bold=True, space_after=Pt(12),
    )

    # ── 10. Signature block ──────────────────────────────────────────────
    section_num += 1
    _add_styled_paragraph(
        doc, f"{section_num}. Подписи", bold=True, space_after=Pt(6))
    _add_styled_paragraph(doc, f"Дата составления: {today}")
    _add_styled_paragraph(doc, "")

    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_headers = ["Должность", "Подпись", "ФИО"]
    for j, h in enumerate(sig_headers):
        _set_cell_text(sig_table.rows[0].cells[j], h, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # Empty row for manual filling
    for j in range(3):
        _set_cell_text(sig_table.rows[1].cells[j], "________________",
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(str(output_path))
    logger.info("Document generated: %s", output_path)
    return str(output_path)


# ── Consolidated Generator ──────────────────────────────────────────────────


def generate_consolidated_document(
    output_dir: Path,
    purchase_id: int,
    region: str | None,
    items: list[dict[str, Any]],
) -> str:
    """Generate a single .docx covering all items in one purchase.

    Each entry in *items* must contain:
        target_name, analogs, prices, outliers, calculation, quantity, unit, justification
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"nmck_consolidated_{purchase_id}.docx"

    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = _FONT_NAME
    style.font.size = _FONT_SIZE

    today = datetime.now().strftime("%d.%m.%Y")

    # ── 1. Title ──
    _add_styled_paragraph(
        doc, "ОБОСНОВАНИЕ", bold=True, size=Pt(14),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_styled_paragraph(
        doc,
        "начальной (максимальной) цены контракта",
        bold=True, size=Pt(14),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(4),
    )
    _add_styled_paragraph(
        doc, "Сводный расчёт по закупке",
        bold=True, size=Pt(12),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(6),
    )
    _add_styled_paragraph(doc, f"Дата: {today}")
    if region:
        _add_styled_paragraph(doc, f"Регион: {region}", space_after=Pt(12))

    # ── 2. Summary table ──
    _add_styled_paragraph(doc, "1. Сводная таблица", bold=True, space_after=Pt(6))

    n_items = len(items)
    tbl = doc.add_table(rows=1 + n_items + 1, cols=6)
    _style_table(tbl)

    headers = ["№", "Наименование", "Кол-во", "Ед. изм.", "НМЦК за ед., руб.", "Итого НМЦК, руб."]
    for j, h in enumerate(headers):
        _set_cell_text(tbl.rows[0].cells[j], h, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

    grand_total = 0.0
    for i, item in enumerate(items):
        row = tbl.rows[1 + i]
        calc = item.get("calculation", {})
        qty = item.get("quantity", 1.0)
        nmck_unit = calc.get("nmck_per_unit", 0)
        total = calc.get("total_nmck", 0)
        grand_total += total

        _set_cell_text(row.cells[0], str(i + 1), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row.cells[1], item.get("target_name", ""))
        qty_str = str(int(qty)) if qty == int(qty) else _format_price(qty)
        _set_cell_text(row.cells[2], qty_str, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row.cells[3], item.get("unit") or "шт", alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row.cells[4], _format_price(nmck_unit), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(row.cells[5], _format_price(total), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # Total row
    total_row = tbl.rows[1 + n_items]
    total_row.cells[0].merge(total_row.cells[4])
    _set_cell_text(total_row.cells[0], "ИТОГО", bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_text(total_row.cells[5], _format_price(grand_total), bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    _add_styled_paragraph(doc, "", space_after=Pt(12))

    # ── 3. Method (once) ──
    _add_styled_paragraph(doc, "2. Метод определения НМЦК", bold=True, space_after=Pt(6))
    _add_styled_paragraph(
        doc,
        "Начальная (максимальная) цена контракта определена "
        "методом сопоставимых рыночных цен (анализа рынка) "
        "в соответствии с ч. 6 ст. 22 Федерального закона от 05.04.2013 "
        "№ 44-ФЗ «О контрактной системе в сфере закупок товаров, работ, "
        "услуг для обеспечения государственных и муниципальных нужд» "
        "и Приказа Минэкономразвития России от 02.10.2013 № 567 "
        "«Об утверждении Методических рекомендаций по применению методов "
        "определения начальной (максимальной) цены контракта, цены "
        "контракта, заключаемого с единственным поставщиком "
        "(подрядчиком, исполнителем)».",
        space_after=Pt(12),
    )

    # ── 4. Per-item details ──
    for i, item in enumerate(items):
        section_base = i + 1
        calc = item.get("calculation", {})
        prices = item.get("prices", [])
        outliers = item.get("outliers", [])
        qty = item.get("quantity", 1.0)
        target_name = item.get("target_name", "")

        doc.add_page_break()
        _add_styled_paragraph(
            doc, f"Позиция {i + 1}: {target_name}",
            bold=True, size=Pt(13),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(12),
        )

        # Sources
        _add_styled_paragraph(doc, "Источники ценовой информации", bold=True, space_after=Pt(6))
        for pi, p in enumerate(prices, 1):
            source = p.get("_source")
            if source == "manual":
                rgn = p.get("Регион заказчика", p.get("region", ""))
                _add_styled_paragraph(doc, f"{pi}. Ручной ввод цены — {rgn}")
            else:
                cid = p.get("Идентификатор контракта", p.get("contract_id", ""))
                date = _format_date(p.get("Дата заключения контракта", p.get("contract_date")))
                rgn = p.get("Регион заказчика", p.get("region", ""))
                _add_styled_paragraph(doc, f"{pi}. Контракт № {cid} от {date}, {rgn}")
        _add_styled_paragraph(doc, "", space_after=Pt(6))

        # Calculation table
        _add_styled_paragraph(doc, "Расчёт", bold=True, space_after=Pt(6))
        n_prices = len(prices)
        table = doc.add_table(rows=1 + n_prices + 6, cols=5)
        _style_table(table)

        tbl_headers = ["№", "Источник", "Цена за ед., руб.", "Врем. коэфф.", "Привед. цена, руб."]
        for j, h in enumerate(tbl_headers):
            _set_cell_text(table.rows[0].cells[j], h, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)

        for pi, p in enumerate(prices):
            row = table.rows[1 + pi]
            raw_price = float(p.get("Цена за единицу", p.get("price", 0)))
            tw = float(p.get("time_weight", 1.0))
            adjusted = raw_price * tw

            source = p.get("_source")
            if source == "manual":
                src_text = "Ручной ввод"
            else:
                cid = p.get("Идентификатор контракта", p.get("contract_id", ""))
                cdate = _format_date(p.get("Дата заключения контракта", p.get("contract_date")))
                src_text = f"Контракт № {cid} от {cdate}"

            _set_cell_text(row.cells[0], str(pi + 1), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row.cells[1], src_text)
            _set_cell_text(row.cells[2], _format_price(raw_price), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell_text(row.cells[3], f"{tw:.4f}", alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row.cells[4], _format_price(adjusted), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

        avg_price = calc.get("weighted_average_price", 0)
        cv = calc.get("coefficient_of_variation", 0)
        nmck_unit = calc.get("nmck_per_unit", 0)
        total_nmck = calc.get("total_nmck", 0)

        summary_data = [
            ("Количество ценовых данных (n)", str(n_prices)),
            ("Средняя цена за единицу, руб.", _format_price(avg_price)),
            ("Коэффициент вариации (КВ), %", f"{cv:.1f}"),
            ("Количество (v)", str(int(qty)) if qty == int(qty) else _format_price(qty)),
            ("НМЦК за единицу, руб.", _format_price(nmck_unit)),
            ("ИТОГО НМЦК, руб.", _format_price(total_nmck)),
        ]
        for idx, (label, value) in enumerate(summary_data):
            r = table.rows[1 + n_prices + idx]
            r.cells[0].merge(r.cells[3])
            _set_cell_text(r.cells[0], label, bold=True)
            _set_cell_text(r.cells[4], value, bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

        # Homogeneity
        is_homo = calc.get("is_homogeneous", True)
        _add_styled_paragraph(doc, "", space_after=Pt(4))
        _add_styled_paragraph(
            doc,
            f"Коэффициент вариации: {cv:.1f}%. "
            f"Совокупность ценовой информации {'однородна' if is_homo else 'неоднородна'}.",
            space_after=Pt(4),
        )

        # Outliers
        if outliers:
            _add_styled_paragraph(
                doc,
                f"Исключено {len(outliers)} позиций как статистически выбивающихся:",
            )
            for o in outliers:
                price_val = float(o.get("Цена за единицу", o.get("price", 0)))
                name_val = o.get("Наименование позиции СТЕ", o.get("cte_name", ""))
                _add_styled_paragraph(doc, f"  — {name_val}: {_format_price(price_val)} руб.")

    # ── 5. Grand total ──
    doc.add_page_break()
    _add_styled_paragraph(doc, "Итого по закупке", bold=True, size=Pt(13),
                          space_after=Pt(12))
    words = number_to_words_ru(grand_total)
    _add_styled_paragraph(
        doc,
        f"Общая НМЦК по закупке составляет: {_format_price(grand_total)} руб. ({words})",
        bold=True, space_after=Pt(12),
    )

    # Signature block
    _add_styled_paragraph(doc, f"Дата составления: {today}", space_after=Pt(6))
    _add_styled_paragraph(doc, "")
    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["Должность", "Подпись", "ФИО"]):
        _set_cell_text(sig_table.rows[0].cells[j], h, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for j in range(3):
        _set_cell_text(sig_table.rows[1].cells[j], "________________",
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(str(output_path))
    logger.info("Consolidated document generated: %s", output_path)
    return str(output_path)
