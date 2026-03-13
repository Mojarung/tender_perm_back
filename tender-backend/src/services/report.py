"""DOCX report generation service using python-docx and Jinja2."""

import io
import logging
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.schemas.api import NMCKReportRequest

logger = logging.getLogger(__name__)


def generate_nmck_report(data: NMCKReportRequest) -> io.BytesIO:
    """
    Generate a .docx report for the NMCC calculation results.
    Returns a BytesIO buffer containing the .docx file.
    """
    doc = Document()

    # Title
    title = doc.add_heading("Расчёт НМЦК", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(
        f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    )
    if data.target_ste_id:
        doc.add_paragraph(f"Идентификатор СТЕ: {data.target_ste_id}")
    if data.target_region:
        doc.add_paragraph(f"Регион заказчика: {data.target_region}")

    doc.add_paragraph("")

    # Results summary
    doc.add_heading("Результаты расчёта", level=2)

    result_table = doc.add_table(rows=3, cols=2)
    result_table.style = "Light Grid Accent 1"
    result_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cells = result_table.rows[0].cells
    cells[0].text = "НМЦК (средняя цена)"
    cells[1].text = f"{data.nmck_value:,.2f} руб."

    cells = result_table.rows[1].cells
    cells[0].text = "Коэффициент вариации"
    cells[1].text = f"{data.variation_coefficient * 100:.2f}%"

    cells = result_table.rows[2].cells
    cells[0].text = "Количество использованных цен"
    cells[1].text = str(len(data.valid_prices_used))

    doc.add_paragraph("")

    # Valid prices table
    doc.add_heading("Использованные цены", level=2)

    prices_table = doc.add_table(rows=1, cols=2)
    prices_table.style = "Light Grid Accent 1"
    header = prices_table.rows[0].cells
    header[0].text = "№"
    header[1].text = "Цена за единицу (руб.)"

    for idx, price in enumerate(data.valid_prices_used, 1):
        row = prices_table.add_row().cells
        row[0].text = str(idx)
        row[1].text = f"{price:,.2f}"

    # Outliers section
    if data.detected_outliers:
        doc.add_paragraph("")
        doc.add_heading("Исключённые цены (выбросы)", level=2)

        outlier_table = doc.add_table(rows=1, cols=2)
        outlier_table.style = "Light Grid Accent 2"
        header = outlier_table.rows[0].cells
        header[0].text = "№"
        header[1].text = "Цена за единицу (руб.)"

        for idx, price in enumerate(data.detected_outliers, 1):
            row = outlier_table.add_row().cells
            row[0].text = str(idx)
            row[1].text = f"{price:,.2f}"

        doc.add_paragraph(
            "Выбросы определены с помощью алгоритма Isolation Forest (sklearn).",
            style="Intense Quote",
        )

    # Methodology note
    doc.add_paragraph("")
    doc.add_heading("Методология", level=2)
    doc.add_paragraph(
        "Расчёт НМЦК выполнен методом анализа рыночных цен в соответствии с "
        "Федеральным законом от 05.04.2013 № 44-ФЗ (ст. 22). "
        "Для определения выбросов применён алгоритм Isolation Forest. "
        "Коэффициент вариации рассчитан как отношение стандартного отклонения "
        "к среднему значению выборки."
    )

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
