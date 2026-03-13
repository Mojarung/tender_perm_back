from docx import Document
import io
import datetime

def generate_nmcc_docx(nmcc_data: dict) -> bytes:
    """
    Generates a .docx file containing the NMCC report calculation based on parameters.
    Returns bytes.
    """
    doc = Document()
    
    # Title
    doc.add_heading('Обоснование Начальной (Максимальной) Цены Контракта (НМЦК)', level=1)
    doc.add_paragraph(f"Дата формирования отчета: {datetime.datetime.now().strftime('%d.%m.%Y')}")
    
    # Target Data
    doc.add_heading('1. Целевая позиция', level=2)
    doc.add_paragraph(f"Идентификатор СТЕ: {nmcc_data.get('target_ste_id', 'Не указан')}")
    doc.add_paragraph(f"Регион поставки: {nmcc_data.get('target_region', 'Не указан')}")
    
    # Mathematical Result
    doc.add_heading('2. Результат расчета', level=2)
    p = doc.add_paragraph()
    p.add_run(f"Итоговая НМЦК (за ед.): {nmcc_data.get('nmck_value', 0.0):.2f}\n").bold = True
    doc.add_paragraph(f"Коэффициент вариации выборки: {nmcc_data.get('variation_coefficient', 0.0):.2f}% (норма: до 33%)")
    
    # AI Explanation
    doc.add_heading('3. Интеллектуальное обоснование (AI)', level=2)
    ai_text = nmcc_data.get('ai_explanation', 'Система не предоставила дополнительного описания.')
    doc.add_paragraph(ai_text)
    
    # Sources Table
    valid_prices = nmcc_data.get('valid_prices_used', [])
    if valid_prices:
        doc.add_heading('4. Реестр использованных источников информации', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№ п/п'
        hdr_cells[1].text = 'Источник (идентификаторы)'
        hdr_cells[2].text = 'Скорректированная цена (руб.)'
        
        for i, price in enumerate(valid_prices, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = "Системный подбор аналогов"
            row_cells[2].text = f"{price:.2f}"
            
    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()
